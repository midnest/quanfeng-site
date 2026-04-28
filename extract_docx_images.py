#!/usr/bin/env python3
"""Extract images and content from Word document"""

import os
import zipfile
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

def extract_docx_images(docx_path, output_dir):
    """Extract images from docx file"""
    os.makedirs(output_dir, exist_ok=True)
    image_count = 0

    # Method 1: Extract from docx as zip (media folder)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            for item in zip_ref.namelist():
                if item.startswith('word/media/'):
                    filename = os.path.basename(item)
                    if filename:
                        output_path = os.path.join(output_dir, f'docx_{filename}')
                        with zip_ref.open(item) as src, open(output_path, 'wb') as dst:
                            dst.write(src.read())
                        image_count += 1
                        print(f"Extracted: {filename}")
    except Exception as e:
        print(f"Error extracting from zip: {e}")

    print(f"\nTotal images extracted: {image_count}")
    return image_count

def analyze_document_structure(docx_path):
    """Analyze document structure"""
    doc = Document(docx_path)

    print("=" * 60)
    print("DOCUMENT STRUCTURE ANALYSIS")
    print("=" * 60)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    # Analyze tables
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables[:10]):  # First 10 tables
        print(f"\nTable {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        # Get first row as header
        if table.rows:
            first_row_text = [cell.text.strip()[:30] for cell in table.rows[0].cells[:5]]
            print(f"  Header: {first_row_text}")

    # Find paragraphs with specific keywords
    print("\n--- KEY PARAGRAPHS ---")
    keywords = ['型号', '规格', '参数', '功率', '电压', 'QA', 'AC', '风机']
    found_sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if any(kw in text for kw in keywords) and len(text) < 100:
            found_sections.append(text)

    # Show unique sections
    unique_sections = list(set(found_sections))[:20]
    for section in unique_sections:
        print(f"  - {section[:60]}")

    return doc

if __name__ == "__main__":
    docx_file = "（母本）交流风机宣传图册（202510更新）(1).docx"
    output_dir = "extracted_docx_images"

    print("Extracting images from Word document...")
    print("-" * 60)

    # Extract images
    count = extract_docx_images(docx_file, output_dir)

    print("\n" + "=" * 60)
    print("Analyzing document structure...")
    print("-" * 60)

    # Analyze structure
    doc = analyze_document_structure(docx_file)

    print(f"\n{'=' * 60}")
    print(f"Images saved to: {output_dir}/")
    print("=" * 60)
